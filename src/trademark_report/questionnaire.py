"""Extract consent fields from the supported software questionnaire variants."""

from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocumentType

from .questionnaire_profiles import (
    QuestionnaireProfile,
    QuestionnaireProfileLibrary,
    default_profile_library,
    normalize_label,
)
from .software_models import (
    QuestionnaireParseResult,
    SoftwareAuthor,
    SoftwareConsentData,
    detect_applicant_type,
)

DATE_RE = re.compile(r"\b(\d{2}\.\d{2}\.\d{4})\s*(?:г\.?|года)?", re.IGNORECASE)
TEXT_DATE_RE = re.compile(
    r"[«\"]?(\d{1,2})[»\"]?\s+"
    r"(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)"
    r"\s+(\d{4})\s*(?:г\.?|года)?",
    re.IGNORECASE,
)
YEAR_FIRST_TEXT_DATE_RE = re.compile(
    r"(\d{4})\s*(?:г\.?|года)?\s+"
    r"[«\"]?(\d{1,2})[»\"]?\s+"
    r"(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)",
    re.IGNORECASE,
)
MONTHS = {
    "января": 1,
    "февраля": 2,
    "марта": 3,
    "апреля": 4,
    "мая": 5,
    "июня": 6,
    "июля": 7,
    "августа": 8,
    "сентября": 9,
    "октября": 10,
    "ноября": 11,
    "декабря": 12,
}
NAME_RE = re.compile(
    r"\b([А-ЯЁ][а-яё-]+(?:-[А-ЯЁ][а-яё-]+)?\s+"
    r"[А-ЯЁ][а-яё-]+(?:-[А-ЯЁ][а-яё-]+)?\s+"
    r"[А-ЯЁ][а-яё-]+(?:-[А-ЯЁ][а-яё-]+)?)\b"
)
MAX_QUESTIONNAIRE_FILE_BYTES = 50 * 1024 * 1024
MAX_QUESTIONNAIRE_UNPACKED_BYTES = 250 * 1024 * 1024
MAX_AUTHORS = 100


def _validate_questionnaire_file(source: Path) -> None:
    """Reject unsupported, damaged and unreasonably large questionnaire files."""

    if source.suffix.casefold() != ".docx":
        raise ValueError("Поддерживаются только анкеты в формате DOCX.")
    try:
        file_size = source.stat().st_size
    except OSError as exc:
        raise ValueError(f"Не удалось открыть анкету: {exc}") from exc
    if file_size > MAX_QUESTIONNAIRE_FILE_BYTES:
        raise ValueError("Анкета слишком большая для безопасной обработки (\u0431олее 50 МБ).")
    try:
        with ZipFile(source) as archive:
            unpacked_size = sum(item.file_size for item in archive.infolist())
    except BadZipFile as exc:
        raise ValueError("Файл не является исправным DOCX-документом.") from exc
    if unpacked_size > MAX_QUESTIONNAIRE_UNPACKED_BYTES:
        raise ValueError("Внутренний размер анкеты слишком велик для безопасной обработки.")


def _clean(value: str) -> str:
    value = value.replace("\xa0", " ").replace("\u200b", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r" *\n *", "\n", value)
    return value.strip(" \n;,")


def _norm(value: str) -> str:
    return normalize_label(value)


def _cell_text(cell) -> str:
    return _clean("\n".join(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip()))


def _rows(
    table,
    library: QuestionnaireProfileLibrary,
    profile: QuestionnaireProfile,
) -> list[tuple[str, str]]:
    result = []
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = _cell_text(row.cells[0])
        value = _cell_text(row.cells[1])
        if library.stops_section(label, profile):
            break
        if label and _norm(label) != _norm(value):
            result.append((label, value))
    return result


def _document_labels(document: DocumentType) -> list[str]:
    labels = []
    for table in document.tables:
        for row in table.rows:
            if row.cells:
                label = _cell_text(row.cells[0])
                if label:
                    labels.append(label)
    return labels


def _section_one_rows(
    document: DocumentType,
    library: QuestionnaireProfileLibrary,
    profile: QuestionnaireProfile,
) -> list[tuple[str, str]]:
    """Collect section-I rows even when Word splits them across several tables."""

    result: list[tuple[str, str]] = []
    started = False
    for table in document.tables:
        whole = _norm(" ".join(cell.text for row in table.rows for cell in row.cells))
        contains_start = library.starts_section(whole, profile)
        row_labels = [
            _cell_text(row.cells[0]) for row in table.rows if row.cells
        ]
        contains_stop_row = any(
            library.stops_section(label, profile) for label in row_labels
        )
        if not started:
            if not contains_start:
                continue
            started = True
        elif contains_start:
            # Some distributed questionnaires append a second, fully filled
            # demonstration form after the client's form. A repeated section
            # start marks a new form, not a continuation of the first one.
            break
        elif row_labels and library.stops_section(row_labels[0], profile):
            break
        result.extend(_rows(table, library, profile))
        if contains_stop_row:
            break
    if not result:
        raise ValueError("В документе не найдена таблица раздела I со сведениями для Роспатента.")
    return result


def _profile_value(
    rows: list[tuple[str, str]],
    library: QuestionnaireProfileLibrary,
    profile: QuestionnaireProfile,
    field_name: str,
) -> str:
    rule = library.field_rule(profile, field_name)
    for label, value in rows:
        if value and rule.matches(label):
            return value
    return ""


def _authors_will_be_mentioned(
    rows: list[tuple[str, str]],
    library: QuestionnaireProfileLibrary,
    profile: QuestionnaireProfile,
) -> bool | None:
    """Read an explicit author-disclosure answer from section I.

    Some questionnaire versions use ``упоминаться``, while others use
    ``указываться``.  An ambiguous, unedited ``да / нет`` value must not hide
    validation errors, so it is represented as ``None``.
    """

    value = _profile_value(rows, library, profile, "author_disclosure")
    normalized = _norm(value)
    if re.match(r"^нет(?:\s|$)", normalized):
        return False
    if re.match(r"^да(?:\s|$)", normalized) and "нет" not in normalized.split():
        return True
    return None


def _extract_inn(value: str) -> str:
    match = re.search(r"(?<!\d)(\d{10}|\d{12})(?!\d)", value.replace(" ", ""))
    return match.group(1) if match else ""


def _extract_applicant(value: str, inn: str) -> tuple[str, str]:
    if not value:
        return "", ""
    text = value
    if inn:
        # Questionnaires often group a personal INN as ``6673 5285 5774``.
        # Match separators horizontally so the expression cannot consume the
        # postal index from the following line.
        spaced_inn = r"[ \t-]*".join(re.escape(digit) for digit in inn)
        text = re.sub(rf"(?<!\d){spaced_inn}(?!\d)", "", text)
    text = re.sub(r"\bИНН\s*[:№-]?", "", text, flags=re.IGNORECASE)
    lines = [_clean(line) for line in text.splitlines() if _clean(line)]
    address = ""
    name_parts = []
    for line in lines:
        postal_index = re.search(r"(?<!\d)\d{6}\s*[,;]", line)
        if postal_index:
            possible_name = line[: postal_index.start()].strip(" ,;-")
            if possible_name:
                name_parts.append(possible_name)
            address = line[postal_index.start() :]
        elif re.search(r"\bадрес\b", line, re.IGNORECASE):
            address = re.sub(r"^.*?адрес\s*[:—-]?\s*", "", line, flags=re.IGNORECASE)
        elif (
            re.search(r"(?i)(?:^|[,;])\s*(?:Россия|РФ)\s*[,;]", line)
            and re.search(r"(?i)\b(?:г\.|город|ул\.|улица|д\.)", line)
        ):
            address = line
        else:
            name_parts.append(line)
    name = _clean(" ".join(name_parts)).strip(" ,;-()")
    return name, address


def _extract_program_name(document: DocumentType) -> str:
    """Read a standalone quoted title used by extended questionnaires."""

    patterns = (
        re.compile(r"(?is)^ПО\s*[«\"]\s*(.+?)\s*[»\"]\s*$"),
        re.compile(r"(?is)^[«\"]\s*(.+?)\s*[»\"]\s*$"),
    )
    for paragraph in document.paragraphs:
        text = _clean(paragraph.text).replace("\n", " ")
        for pattern in patterns:
            match = pattern.match(text)
            if match:
                return _clean(match.group(1))
    return ""


def _extract_names(value: str) -> list[str]:
    result = []
    exact_name = re.compile(
        r"^[А-ЯЁ][А-ЯЁа-яё-]+\s+[А-ЯЁ][А-ЯЁа-яё-]+\s+[А-ЯЁ][А-ЯЁа-яё-]+$"
    )
    for line in value.splitlines():
        candidate = re.sub(r"^\s*[-–—•●]+\s*", "", line).strip(" ,;:.")
        if not exact_name.fullmatch(candidate):
            continue
        if any(word in candidate.upper().split() for word in ("ВЫДАН", "ПАСПОРТ", "РОССИИ")):
            continue
        if candidate.isupper():
            candidate = " ".join(
                "-".join(part.capitalize() for part in word.split("-"))
                for word in candidate.split()
            )
        if candidate not in result:
            result.append(candidate)
    for match in NAME_RE.finditer(value):
        name = match.group(1)
        if name not in result and not any(
            phrase in name for phrase in ("Российская Федерация", "Федерального закона")
        ):
            result.append(name)
    return result


def _normalize_numeric_dates(value: str) -> str:
    def normalized_year(raw_year: str) -> int:
        year = int(raw_year)
        if len(raw_year) == 4:
            return year
        current_short_year = date.today().year % 100
        return 2000 + year if year <= current_short_year else 1900 + year

    return re.sub(
        r"(?<!\d)(\d{1,2})\.\s*(\d{1,2})\.\s*(\d{2}|\d{4})(?!\d)(?!\.\d)",
        lambda match: (
            f"{int(match.group(1)):02d}.{int(match.group(2)):02d}."
            f"{normalized_year(match.group(3)):04d}"
        ),
        value,
    )


def _text_date_value(match: re.Match) -> str:
    return f"{int(match.group(1)):02d}.{MONTHS[match.group(2).lower()]:02d}.{match.group(3)}"


def _year_first_text_date_value(match: re.Match) -> str:
    return f"{int(match.group(2)):02d}.{MONTHS[match.group(3).lower()]:02d}.{match.group(1)}"


def _looks_like_passport_value(value: str) -> bool:
    text = _normalize_numeric_dates(value)
    has_identity = bool(_extract_names(text)) and bool(DATE_RE.search(text) or TEXT_DATE_RE.search(text))
    has_passport = bool(
        re.search(r"(?i)\bпаспорт\b", text)
        or re.search(
            r"(?<!\d)(?:\d{2}[ \t]+\d{2}|\d{4})[ \t]+\d{6}(?!\d)",
            text,
        )
        or re.search(r"(?i)\b[А-ЯЁA-Z]{2}\s*№\s*\d{6,9}\b", text)
    )
    return has_identity and has_passport


def _passport_only_text(value: str) -> str:
    lines = []
    for line in value.splitlines():
        if re.match(
            r"(?i)^\s*[-–—•●]+\s*(?:разработ|написан|системн|формализац|"
            r"программ|тестирован|проектирован|создан|проработ|реализац)",
            line,
        ):
            break
        lines.append(line)
    return "\n".join(lines)


def _split_author_chunks(value: str) -> list[str]:
    text = _clean(value)
    if not text:
        return []
    passport_heading = re.compile(
        r"(?im)(?=^[ \t]*паспортные\s+данные\s+автора?\s+\d+\s*\*?\s*$)"
    )
    passport_heading_starts = [match.start() for match in passport_heading.finditer(text)]
    if passport_heading_starts:
        return [chunk for chunk in [
            _clean(
                text[
                    start : passport_heading_starts[index + 1]
                    if index + 1 < len(passport_heading_starts)
                    else None
                ]
            )
            for index, start in enumerate(passport_heading_starts)
        ] if chunk]
    explicit = re.compile(r"(?im)(?=^[ \t]*Автор[ \t]+\d+[ \t]*:)")
    explicit_starts = [match.start() for match in explicit.finditer(text)]
    if explicit_starts:
        return [chunk for chunk in [
            _clean(text[start : explicit_starts[index + 1] if index + 1 < len(explicit_starts) else None])
            for index, start in enumerate(explicit_starts)
        ] if chunk]
    marker = re.compile(
        r"(?im)(?:^|\n)\s*(?:автор\s*)?\d+\s*[.:)]\s*"
        r"(?=[А-ЯЁ][а-яё-]+\s+[А-ЯЁ][а-яё-]+\s+[А-ЯЁ][а-яё-]+)"
    )
    matches = list(marker.finditer(text))
    if not matches:
        return [text]
    chunks = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        chunks.append(_clean(text[match.end() : end]))
    return [chunk for chunk in chunks if chunk]


def _strip_author_prefix(value: str) -> str:
    return _clean(re.sub(r"(?i)^\s*(?:автор\s*)?\d+\s*[.:)]\s*", "", value))


def _clean_passport_issuer(value: str) -> str:
    """Remove questionnaire-only fields accidentally captured as the issuer."""

    value = re.sub(
        r"(?i)^\s*(?:года\b|г\.(?=\s|$)|г(?=\s))\s*[,;:-]?\s*",
        "",
        value,
    )
    value = re.sub(
        r"(?i)^\s*(?:выдавший\s+орган|орган,?\s+выдавший\s+документ)\s*"
        r"[:—-]?\s*(?:орган\s+)?",
        "",
        value,
    )
    value = re.split(
        r"(?i)\b(?:код\s+подразделения|к\.?\s*п\.?|дата\s+выдачи|"
        r"зарегистрирован\w*|проживающ\w*|место\s+жительства)\b\s*[:—-]?",
        value,
        maxsplit=1,
    )[0]
    return re.sub(r"\s+", " ", value.replace("\u200b", "")).strip(" ,;:-")


def _clean_author_address(value: str) -> str:
    value = re.sub(
        r"(?i)^\s*(?:(?:адрес\s+)?места?\s+жительства|"
        r"зарегистрирован\w*|проживающ\w*)\s*[:—-]?\s*",
        "",
        value,
    )
    value = re.sub(
        r"(?i)\s*,?\s+от\s+\d{2}\.\d{2}\.\d{4}\s*(?:г\.?|года)?\s*\.?$",
        "",
        value,
    )
    return _clean(value).lstrip("-–—•● ").strip(" ;")


def _extract_passport(author: SoftwareAuthor, value: str) -> None:
    text = _normalize_numeric_dates(_strip_author_prefix(_passport_only_text(value)))
    author.source_text = text
    names = _extract_names(text)
    if names and not author.full_name:
        author.full_name = names[0]

    birth = re.search(r"(?i)(?:дата\s+рождения\s*[:—-]?\s*|)(\d{2}\.\d{2}\.\d{4})\s*(?:г\.?р\.?)", text)
    if not birth:
        birth = re.search(r"(?i)дата\s+рождения\s*[:—-]?\s*(\d{2}\.\d{2}\.\d{4})", text)
    year_first_birth = YEAR_FIRST_TEXT_DATE_RE.search(text)
    dates = DATE_RE.findall(text)
    if birth:
        author.birth_date = birth.group(1)
    elif year_first_birth:
        author.birth_date = _year_first_text_date_value(year_first_birth)
    elif dates:
        author.birth_date = dates[0]

    passport_patterns = (
        r"(?i)паспорт(?:\s+гражданина\s+РФ)?\s*[:—-]?\s*(?:серия|серии)?\s*[:—-]?\s*[«\"]?(\d{2})\s*(\d{2})[»\"]?\s*[,;]?\s*(?:номер|№)?\s*[«\"]?(\d{6})",
        r"(?i)паспорт(?:\s+РФ)?\s*(\d{4})\s*(\d{6})",
        r"(?i)(?:серия|серии)\s*[:—-]?\s*[«\"]?(\d{2})\s*(\d{2})[»\"]?\s*[,;]?\s*(?:и\s+)?(?:номер|№|no\.?)\s*[«\"]?(\d{6})",
        r"(?i)паспорт\s+([А-ЯЁA-Z]{2})\s*№\s*(\d{6,9})",
        r"(?<!\d)(\d{2})[ \t]+(\d{2})[ \t]+(\d{6})(?!\d)",
        r"(?<!\d)(\d{4})[ \t]+(\d{6})(?!\d)",
    )
    passport_end: int | None = None
    for pattern in passport_patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        groups = match.groups()
        if len(groups) == 3:
            author.passport_series = groups[0] + groups[1]
            author.passport_number = groups[2]
        else:
            author.passport_series, author.passport_number = groups
        passport_end = match.end()
        break

    indexed = {
        int(n): _clean(v)
        for n, v in re.findall(r"(?m)^[ \t]*([1-5])\.[ \t]*(.+)$", text)
    }
    if indexed:
        author.full_name = author.full_name or indexed.get(1, "")
        indexed_birth = DATE_RE.search(indexed.get(2, ""))
        author.birth_date = author.birth_date or (indexed_birth.group(1) if indexed_birth else "")
        author.address = author.address or indexed.get(3, "")
        digits = re.sub(r"\D", "", indexed.get(4, ""))
        if len(digits) >= 10:
            author.passport_series = author.passport_series or digits[:4]
            author.passport_number = author.passport_number or digits[4:10]
        author.passport_issuer = author.passport_issuer or indexed.get(5, "")
        issue_dates = DATE_RE.findall(indexed.get(5, ""))
        if issue_dates:
            author.passport_issue_date = issue_dates[-1]
            author.passport_issuer = _clean(indexed[5].replace(issue_dates[-1], ""))

    issue_label = re.search(r"(?i)дата\s+выдачи\s*[:—-]?\s*(\d{2}\.\d{2}\.\d{4})", text)
    issued_before = re.search(
        r"(?is)\bвыдан(?:ный|а)?\s*[:—-]?\s*(\d{2}\.\d{2}\.\d{4})\s*(?:г\.\s*)?(.+?)(?=\b(?:зарегистр|прожива|место\s+жительства|дата\s+рождения|код\s+подразделения)\b|$)",
        text,
    )
    issued_after = re.search(
        r"(?is)\bвыдан(?:ный|а)?\s*[:—-]?\s*(.+?)\s+"
        r"(\d{2}\.\d{2}\.\d{4})(?=\s*[,;.]|\s|$)",
        text,
    )
    text_date_match = TEXT_DATE_RE.search(text)
    text_issue_date = _text_date_value(text_date_match) if text_date_match else ""
    text_date_issuer = ""
    if text_date_match:
        line_start = text.rfind("\n", 0, text_date_match.start()) + 1
        line_end = text.find("\n", text_date_match.end())
        if line_end < 0:
            line_end = len(text)
        line = text[line_start:line_end]
        marker = re.search(r"(?i)\bвыдан(?:ный|а)?\b", line)
        if marker:
            relative_date_start = text_date_match.start() - line_start
            before = line[: marker.start()].strip(" ,;:-")
            after_marker = line[marker.end() : relative_date_start].strip(" ,;:-")
            after_date = line[text_date_match.end() - line_start :].strip(" ,;:-")
            if before and not re.search(r"(?i)\bпаспорт\b|\d{6}", before):
                text_date_issuer = before
            elif after_marker:
                text_date_issuer = after_marker
            elif after_date:
                text_date_issuer = after_date
    if issue_label:
        author.passport_issue_date = issue_label.group(1)
    elif issued_before:
        author.passport_issue_date = issued_before.group(1)
    elif issued_after:
        author.passport_issue_date = issued_after.group(2)
    elif text_issue_date:
        author.passport_issue_date = text_issue_date
    elif len(dates) >= 2:
        author.passport_issue_date = next((item for item in dates if item != author.birth_date), "")

    if issued_before:
        author.passport_issuer = _clean(issued_before.group(2))
    elif issued_after:
        author.passport_issuer = _clean(issued_after.group(1))
    elif text_date_issuer:
        author.passport_issuer = _clean(text_date_issuer)
    else:
        issuer = re.search(
            r"(?is)\bвыдан(?:ный|а)?\s*[:—-]?\s*(.+?)(?=\b(?:код\s+подразделения|дата\s+выдачи|место\s+жительства|дата\s+рождения|зарегистр)\b|$)",
            text,
        )
        if issuer:
            author.passport_issuer = _clean(issuer.group(1))
    if not author.passport_issuer and author.passport_issue_date:
        date_position = text.find(author.passport_issue_date)
        if date_position >= 0:
            possible_issuer = text[date_position + len(author.passport_issue_date) :]
            possible_issuer = possible_issuer.split("\n", 1)[0].strip(" ,;:-")
            if (
                possible_issuer
                and not re.fullmatch(r"\d{3}\s*[-–—]\s*\d{3}", possible_issuer)
                and not re.match(
                    r"(?i)^(?:зарегистрирован\w*|прописка|проживающ\w*|адрес)\b",
                    possible_issuer,
                )
            ):
                author.passport_issuer = _clean(possible_issuer)
            else:
                line_start = text.rfind("\n", 0, date_position) + 1
                before_date = text[line_start:date_position]
                passport_pair = re.search(
                    r"(?<!\d)(?:\d{2}[ \t]+\d{2}|\d{4})[ \t]+\d{6}(?!\d)",
                    before_date,
                )
                if passport_pair:
                    possible_issuer = before_date[passport_pair.end() :].strip(" ,;:-")
                    if possible_issuer:
                        author.passport_issuer = _clean(possible_issuer)
    if not author.passport_issuer:
        issuer = re.search(
            r"(?is)\b(?:выдавший\s+орган|орган\s+выдачи)\s*[:–—-]\s*"
            r"(.+?)(?=\b(?:зарегистр|прожива|место\s+жительства|"
            r"код\s+подразделения)\b|$)",
            text,
        )
        if issuer:
            author.passport_issuer = _clean(issuer.group(1))
    if not author.passport_issuer:
        issuer = re.search(
            r"(?is)орган,?\s+выдавший\s+документ\s*:\s*(?:орган\s+)?(.+?)(?=\b(?:зарегистр|прожива)\b|$)",
            text,
        )
        if issuer:
            author.passport_issuer = _clean(issuer.group(1))
    if author.passport_issue_date and author.passport_issuer:
        author.passport_issuer = _clean(author.passport_issuer.replace(author.passport_issue_date, ""))
    if not author.passport_issuer and passport_end is not None and author.passport_issue_date:
        date_position = text.find(author.passport_issue_date, passport_end)
        if date_position >= 0:
            between_passport_and_date = text[passport_end:date_position].strip(" ,;:-–—")
            if re.search(r"[А-ЯЁа-яёA-Za-z]", between_passport_and_date):
                author.passport_issuer = _clean(between_passport_and_date)
    author.passport_issuer = _clean_passport_issuer(author.passport_issuer)

    address_patterns = (
        r"(?is)(?:зарегистрирован\w*|проживающ\w*|прописка|адрес)"
        r"(?:\s+(?:по\s+)?адресу)?\s*[:—-]?\s*"
        r"(.+?)(?=\n\s*[-–—•●]*\s*(?:серия|паспорт|дата\s+выдачи|выдан|"
        r"дата\s+рождения|разработ|написан|системн|формализац|программ|"
        r"тестирован)|$)",
        r"(?is)место\s+жительства\s*[:—-]?\s*(.+?)"
        r"(?=\n\s*[-–—•●]*\s*(?:серия|паспорт|дата\s+выдачи|"
        r"дата\s+рождения|выдан)|$)",
    )
    for pattern in address_patterns:
        match = re.search(pattern, text)
        if match:
            author.address = _clean_author_address(match.group(1))
            break
    if not author.address:
        lines = [_clean(line) for line in text.splitlines() if _clean(line)]
        for line_index, line in enumerate(lines):
            cleaned_line = re.sub(r"^\s*[-–—•●]+\s*", "", line).strip(" ;")
            date_only = re.fullmatch(
                r"(?i)\d{2}\.\d{2}\.\d{4}\s*(?:г\.?р\.?|года\s+рождения)?",
                cleaned_line,
            )
            textual_date_only = bool(
                TEXT_DATE_RE.fullmatch(cleaned_line)
                or YEAR_FIRST_TEXT_DATE_RE.fullmatch(cleaned_line)
            )
            passport_numbers = re.search(
                r"(?<!\d)(?:\d{2}[ \t]+\d{2}|\d{4})[ \t]+\d{6}(?!\d)",
                cleaned_line,
            )
            if (
                re.search(
                    r"(?i)(?:\b(?:россия|российская\s+федерация|обл\.?|край|город|ул\.?|дом)\b|"
                    r"\bг\.|\bд\.|(?<!\d)\d{6}\s*,)",
                    cleaned_line,
                )
                and not re.search(r"(?i)паспорт|выдан|дата\s+рождения", cleaned_line)
                and not date_only
                and not textual_date_only
                and not passport_numbers
            ):
                if author.full_name not in line:
                    address_lines = [cleaned_line]
                    for continuation in lines[line_index + 1 :]:
                        continuation = re.sub(
                            r"^\s*[-–—•●]+\s*", "", continuation
                        ).strip(" ;")
                        if re.search(
                            r"(?i)\b(?:паспорт|серия|номер|выдан|"
                            r"дата\s+выдачи|код\s+подразделения)\b",
                            continuation,
                        ):
                            break
                        if not (
                            address_lines[-1].rstrip().endswith(",")
                            or re.match(
                                r"(?i)^(?:ул\.?|улица|проспект|пр-т|пер\.?|"
                                r"переулок|шоссе|наб\.?|набережная|д\.?|дом|"
                                r"корп\.?|корпус|стр\.?|строение|кв\.?|квартира)\b",
                                continuation,
                            )
                        ):
                            break
                        address_lines.append(continuation)
                    author.address = _clean_author_address(", ".join(address_lines))
                    break


def _clean_contribution(value: str) -> str:
    value = re.sub(r"(?i)^\s*Автор\s+\d+\s*:\s*", "", value)
    lines = []
    for line in value.splitlines():
        line = re.sub(r"^\s*[-–—•●]+\s*", "", line).strip()
        if line:
            lines.append(line.rstrip(";."))
    if len(lines) > 1 and NAME_RE.fullmatch(lines[0].rstrip(".")):
        lines.pop(0)
    return "; ".join(lines)


def _contribution_from_author_chunk(value: str) -> str:
    """Extract bullet-point contribution lines belonging to one grouped author."""

    lines = [
        line
        for line in value.splitlines()
        if re.match(
            r"(?i)^\s*[-–—•●]+\s*(?:разработ|написан|системн|формализац|"
            r"программ|тестирован|проектирован|создан|проработ|реализац)",
            line,
        )
    ]
    return _clean_contribution("\n".join(lines))


def _named_contributions(value: str) -> dict[str, str]:
    """Extract contributions keyed by FIO instead of relying on row order."""

    header = re.compile(
        r"(?im)^[ \t]*(?:\d+[.)][ \t]*)?"
        r"([А-ЯЁ][А-ЯЁа-яё-]+(?:-[А-ЯЁ][А-ЯЁа-яё-]+)?[ \t]+"
        r"[А-ЯЁ][А-ЯЁа-яё-]+(?:-[А-ЯЁ][А-ЯЁа-яё-]+)?[ \t]+"
        r"[А-ЯЁ][А-ЯЁа-яё-]+(?:-[А-ЯЁ][А-ЯЁа-яё-]+)?)"
        r"[ \t]*[-–—:][ \t]*"
    )
    matches = list(header.finditer(value))
    result: dict[str, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(value)
        contribution = _clean_contribution(value[match.end() : end])
        if contribution:
            result[_norm(match.group(1))] = contribution
    return result


def parse_questionnaire(
    path: str | Path,
    *,
    profile_library: QuestionnaireProfileLibrary | None = None,
) -> QuestionnaireParseResult:
    source = Path(path)
    _validate_questionnaire_file(source)
    document = Document(str(source))
    library = profile_library or default_profile_library()
    profile = library.select_profile(_document_labels(document))
    rows = _section_one_rows(document, library, profile)
    warnings: list[str] = []
    authors_will_be_mentioned = _authors_will_be_mentioned(rows, library, profile)

    program_name = _profile_value(
        rows, library, profile, "program_name"
    ) or _extract_program_name(document)
    applicant_raw = _profile_value(rows, library, profile, "applicant")
    inn = _extract_inn(applicant_raw)
    applicant_name, applicant_address = _extract_applicant(applicant_raw, inn)
    count_value = _profile_value(rows, library, profile, "author_count")
    count_match = re.search(r"\d+", count_value)
    declared_count = int(count_match.group()) if count_match else None

    names_value = _profile_value(rows, library, profile, "author_names")
    known_names = _extract_names(names_value)
    authors: list[SoftwareAuthor] = [SoftwareAuthor(full_name=name) for name in known_names]
    by_number: defaultdict[int, dict[str, list[str]]] = defaultdict(lambda: defaultdict(list))
    sequential_passports: list[str] = []
    sequential_contributions: list[str] = []
    contributions_by_name: dict[str, str] = {}
    sequential_bases: list[str] = []
    grouped_passport_count = 0

    for label, value in rows:
        normalized = _norm(label)
        if not value:
            continue
        number_match = re.search(r"автора?\s*(\d+)", normalized)
        number = int(number_match.group(1)) if number_match else None
        looks_passport = library.row_matches(profile, "passport", normalized)
        looks_contribution = library.row_matches(profile, "contribution", normalized)
        looks_basis = library.row_matches(profile, "rights_basis", normalized)
        named_contributions = (
            _named_contributions(value) if looks_contribution else {}
        )
        contributions_by_name.update(named_contributions)

        if looks_basis:
            chunks = _split_author_chunks(value)
            if number:
                by_number[number]["basis"].append(value)
            elif len(chunks) > 1:
                sequential_bases.extend(chunks)
            elif grouped_passport_count > 1:
                sequential_bases = [value] * grouped_passport_count
            else:
                sequential_bases.append(value)
            continue
        if looks_passport and looks_contribution:
            if _looks_like_passport_value(value) or re.search(
                r"(?i)паспорт|серия|номер|выдан|дата рождения|г\.р\.", value
            ):
                chunks = _split_author_chunks(value)
                if len(chunks) > 1:
                    sequential_passports.extend(chunks)
                    grouped_passport_count = max(grouped_passport_count, len(chunks))
                    sequential_contributions.extend(
                        _contribution_from_author_chunk(chunk) for chunk in chunks
                    )
                    continue
                else:
                    (by_number[number]["passport"] if number else sequential_passports).append(value)
                if re.search(r"(?im)^\s*[-–—•●]+\s*(?:разработ|написан|системн|формализац|программ)", value):
                    contribution = "\n".join(
                        line for line in value.splitlines()
                        if re.match(r"(?i)^\s*[-–—•●]+\s*(?:разработ|написан|системн|формализац|программ)", line)
                    )
                    if contribution:
                        (by_number[number]["contribution"] if number else sequential_contributions).append(contribution)
            else:
                if named_contributions:
                    continue
                chunks = _split_author_chunks(value)
                if len(chunks) > 1:
                    if (
                        grouped_passport_count == len(chunks)
                        and len(sequential_contributions) == grouped_passport_count
                    ):
                        for chunk_index, chunk in enumerate(chunks):
                            if not sequential_contributions[chunk_index].strip():
                                sequential_contributions[chunk_index] = chunk
                    else:
                        sequential_contributions.extend(chunks)
                elif grouped_passport_count > 1:
                    if not any(item.strip() for item in sequential_contributions):
                        sequential_contributions = [value] * grouped_passport_count
                else:
                    (by_number[number]["contribution"] if number else sequential_contributions).append(value)
            continue
        if looks_passport:
            chunks = _split_author_chunks(value)
            if len(chunks) > 1:
                sequential_passports.extend(chunks)
                grouped_passport_count = max(grouped_passport_count, len(chunks))
            elif number:
                by_number[number]["passport"].append(value)
            elif chunks:
                sequential_passports.extend(chunks)
            continue
        if looks_contribution:
            chunks = _split_author_chunks(value)
            if len(chunks) > 1:
                sequential_contributions.extend(chunks)
            elif number:
                by_number[number]["contribution"].append(value)
            elif chunks:
                sequential_contributions.extend(chunks)

    target_count = max(
        declared_count or 0,
        len(authors),
        max(by_number, default=0),
        len(sequential_passports),
        len(sequential_contributions),
    )
    if target_count > MAX_AUTHORS:
        raise ValueError(
            f"В анкете распознано неожиданно большое число авторов ({target_count}). "
            "Проверьте поле с количеством авторов."
        )
    while len(authors) < target_count:
        authors.append(SoftwareAuthor())

    for index, author in enumerate(authors, 1):
        passport_values = by_number[index]["passport"] or (
            [sequential_passports[index - 1]] if index <= len(sequential_passports) else []
        )
        if passport_values:
            _extract_passport(author, "\n".join(passport_values))
        named_contribution = contributions_by_name.get(_norm(author.full_name), "")
        contribution_values = (
            by_number[index]["contribution"]
            or ([named_contribution] if named_contribution else [])
            or (
                [sequential_contributions[index - 1]]
                if index <= len(sequential_contributions)
                else []
            )
        )
        if contribution_values:
            author.creative_contribution = _clean_contribution("\n".join(contribution_values))
        basis_values = by_number[index]["basis"] or (
            [sequential_bases[index - 1]] if index <= len(sequential_bases) else []
        )
        if basis_values:
            author.rights_basis = _clean_contribution("\n".join(basis_values))

    if declared_count is not None and len(authors) != declared_count:
        warnings.append(
            f"В анкете указано авторов: {declared_count}; распознано карточек: {len(authors)}."
        )
    if not authors:
        warnings.append("В разделе I не найдены сведения об авторах.")
    for index, author in enumerate(authors, 1):
        missing = author.missing_fields(
            require_personal_data=authors_will_be_mentioned is not False
        )
        if missing:
            warnings.append(f"Автор {index}: не найдены {', '.join(missing)}.")
    if not applicant_address:
        warnings.append("В разделе I не найден адрес заявителя — его нужно заполнить вручную.")

    data = SoftwareConsentData(
        program_name=program_name,
        applicant_name=applicant_name,
        applicant_address=applicant_address,
        inn=inn,
        applicant_type=detect_applicant_type(applicant_name, inn),
        authors=authors,
        authors_will_be_mentioned=authors_will_be_mentioned,
        declared_author_count=declared_count,
        source_path=str(source),
        questionnaire_profile_id=profile.profile_id,
        questionnaire_profile_name=profile.name,
    )
    if not data.program_name:
        warnings.insert(0, "В разделе I не найдено название программы.")
    if not data.applicant_name:
        warnings.insert(0, "В разделе I не найдено наименование заявителя.")
    if not data.inn:
        warnings.insert(0, "В разделе I не найден ИНН заявителя.")
    return QuestionnaireParseResult(data=data, warnings=warnings)
