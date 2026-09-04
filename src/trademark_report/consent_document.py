"""Create a combined, two-page-per-author consent DOCX from the retained template."""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from dataclasses import replace
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from .docx_utils import strip_comments
from .io_utils import atomic_output_path, atomic_write_bytes
from .software_models import (
    APPLICANT_INDIVIDUAL,
    APPLICANT_SOLE_PROPRIETOR,
    SoftwareAuthor,
    SoftwareConsentData,
)


def _resource_root() -> Path:
    bundle_root = getattr(sys, "_MEIPASS", None)
    if getattr(sys, "frozen", False) and bundle_root:
        return Path(bundle_root)
    return Path(__file__).resolve().parents[2]


DEFAULT_CONSENT_TEMPLATE = _resource_root() / "assets" / "software_consent_template.docx"
TABLE_WIDTH_EXTENSION = Cm(0.5)


def _iter_paragraphs(document: DocumentType):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs


def _replace_in_runs(paragraph, old: str, new: str, *, count: int = -1) -> int:
    replaced = 0
    while count < 0 or replaced < count:
        runs = paragraph.runs
        full = "".join(run.text for run in runs)
        start = full.find(old)
        if start < 0:
            break
        end = start + len(old)
        positions = []
        cursor = 0
        for index, run in enumerate(runs):
            positions.append((index, cursor, cursor + len(run.text)))
            cursor += len(run.text)
        first = next((item for item in positions if item[1] <= start < item[2] or item[1] == start == item[2]), None)
        last = next((item for item in reversed(positions) if item[1] < end <= item[2]), None)
        if first is None or last is None:
            break
        first_index, first_start, _ = first
        last_index, last_start, _ = last
        prefix = runs[first_index].text[: start - first_start]
        suffix = runs[last_index].text[end - last_start :]
        runs[first_index].text = prefix + new + suffix
        for index in range(first_index + 1, last_index + 1):
            runs[index].text = ""
        replaced += 1
    return replaced


def _set_text_preserve(paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _expand_table_width(table, extra_width=TABLE_WIDTH_EXTENSION) -> None:
    """Expand a fixed-width template table while preserving column proportions."""

    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None or table_width.get(qn("w:type")) != "dxa":
        return
    old_width = int(table_width.get(qn("w:w"), "0"))
    if old_width <= 0:
        return
    new_width = old_width + extra_width.twips
    table_width.set(qn("w:w"), str(new_width))

    grid_columns = list(table._tbl.tblGrid)
    old_grid_widths = [int(column.get(qn("w:w"), "0")) for column in grid_columns]
    if old_grid_widths and sum(old_grid_widths) > 0:
        remaining = new_width
        for index, (column, width) in enumerate(
            zip(grid_columns, old_grid_widths, strict=True)
        ):
            scaled = (
                remaining
                if index == len(grid_columns) - 1
                else round(width * new_width / old_width)
            )
            column.set(qn("w:w"), str(scaled))
            remaining -= scaled

    for cell in table._tbl.iter(qn("w:tc")):
        cell_width = cell.find(f"{qn('w:tcPr')}/{qn('w:tcW')}")
        if cell_width is None or cell_width.get(qn("w:type")) != "dxa":
            continue
        width = int(cell_width.get(qn("w:w"), "0"))
        cell_width.set(qn("w:w"), str(round(width * new_width / old_width)))


def _passport_text(author: SoftwareAuthor) -> str:
    def single_line(value: str) -> str:
        return re.sub(r"\s+", " ", value.replace("\u200b", "")).strip()

    issuer = re.split(
        r"(?i)\b(?:код\s+подразделения|дата\s+выдачи)\s*[:—-]?",
        author.passport_issuer,
        maxsplit=1,
    )[0]
    issuer = single_line(issuer).strip(" ,;:-")
    when_and_where = " ".join(
        item for item in (single_line(author.passport_issue_date), issuer) if item
    )
    return (
        f"Паспорт гражданина РФ, серия «{single_line(author.passport_series)}» "
        f"номер «{single_line(author.passport_number)}» выдан {when_and_where}"
    )


def _fill_author_document(
    document: DocumentType,
    data: SoftwareConsentData,
    author: SoftwareAuthor,
) -> None:
    date_text = data.document_date.strftime("%d.%m.%Y")
    birth = None
    try:
        birth = datetime.strptime(author.birth_date, "%d.%m.%Y")
    except ValueError:
        pass
    page_one = True
    exact_ellipsis_seen = 0
    signature_seen = 0
    date_seen = 0

    for paragraph in _iter_paragraphs(document):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("Согласие автора на указание сведений"):
            page_one = False
        if "Название программы для ЭВМ или базы данных" in text:
            continue
        if text == "«…»":
            exact_ellipsis_seen += 1
            is_program_name = page_one or exact_ellipsis_seen == 1
            value = data.program_name if is_program_name else author.creative_contribution
            _set_text_preserve(paragraph, f"«{value}»" if is_program_name else value)
            if not page_one and exact_ellipsis_seen > 1:
                size = 7.5 if len(value) > 500 else 8.5 if len(value) > 300 else 9.5
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(size)
            continue
        if text.startswith("№ заявки"):
            _remove_paragraph(paragraph)
            continue
        if text.startswith("(указывается при наличии регистрационного номера заявки"):
            _remove_paragraph(paragraph)
            continue
        if text.startswith("Заявка №"):
            _remove_paragraph(paragraph)
            continue
        if text.startswith("Название:"):
            _set_text_preserve(paragraph, f"Название: «{data.program_name}»")
            continue
        if text.startswith("Ф. И. О. субъекта персональных данных"):
            _set_text_preserve(paragraph, f"Ф. И. О. субъекта персональных данных  {author.full_name}")
            continue
        if text.startswith("Адрес места жительства"):
            _set_text_preserve(paragraph, f"Адрес места жительства  {author.address}")
            continue
        if text.startswith("Документ, удостоверяющий личность"):
            prefix = "Документ, удостоверяющий личность субъекта персональных данных, дата его выдачи и выдавший орган  "
            _set_text_preserve(paragraph, prefix + _passport_text(author))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue
        if text.startswith("Общество с ограниченной ответственностью"):
            _set_text_preserve(paragraph, data.applicant_name)
            continue
        if text.startswith("Россия, «") and text.endswith("»"):
            _set_text_preserve(paragraph, data.applicant_address)
            continue
        if text.startswith("ОГРН:"):
            if data.applicant_type == APPLICANT_INDIVIDUAL:
                registration_text = f"ИНН: {data.inn}"
            elif data.applicant_type == APPLICANT_SOLE_PROPRIETOR:
                registration_text = f"ОГРНИП: {data.ogrn}     ИНН: {data.inn}"
            else:
                registration_text = f"ОГРН: {data.ogrn}     ИНН: {data.inn}"
            _set_text_preserve(paragraph, registration_text)
            continue
        if text.startswith("Фамилия имя отчество:"):
            _set_text_preserve(paragraph, f"Фамилия имя отчество: {author.full_name}")
            continue
        if text.startswith("Дата рождения:"):
            if birth:
                value = (
                    f"Дата рождения: число: {birth:%d}   месяц: {birth:%m}   год: {birth:%Y}   "
                    f"Гражданство: {author.citizenship}"
                )
            else:
                value = f"Дата рождения: {author.birth_date}   Гражданство: {author.citizenship}"
            _set_text_preserve(paragraph, value)
            continue
        if (
            data.authors_will_be_mentioned is False
            and "упоминать его под своим именем" in text
            and "не упоминать его (анонимно)" in text
        ):
            for run in paragraph.runs:
                run.text = (
                    run.text.replace("☒", "__CHECKED_BOX__")
                    .replace("☐", "☒")
                    .replace("__CHECKED_BOX__", "☐")
                )
            continue
        if text.startswith("Россия, «") and text.endswith("» (RU)"):
            _set_text_preserve(paragraph, f"{author.address} (RU)")
            continue
        if text.startswith("Подпись") and "ФИО" in text:
            signature_seen += 1
            label = "Подпись автора:" if not page_one else "Подпись"
            if not page_one:
                paragraph.insert_paragraph_before(style=paragraph.style)
            _set_text_preserve(paragraph, f"{label} ___________________/ {author.full_name} /")
            continue
        if text == "Дата":
            date_seen += 1
            _set_text_preserve(paragraph, date_text)

    # The template reserves a very tall final row for the attorney/date block.
    # A smaller minimum keeps long, unabridged author contributions on page 2.
    if document.tables and len(document.tables[0].rows) >= 6:
        _expand_table_width(document.tables[0])
        # Reclaim a small amount of unused vertical padding in the title row.
        # This keeps the form on two pages after enlarging both signature areas.
        title_row_properties = document.tables[0].rows[2]._tr.get_or_add_trPr()
        title_row_height = title_row_properties.find(qn("w:trHeight"))
        if title_row_height is None:
            title_row_height = OxmlElement("w:trHeight")
            title_row_properties.append(title_row_height)
        title_row_height.set(qn("w:val"), "900")

        row = document.tables[0].rows[5]
        cell = row.cells[0]
        empty_after_text = False
        seen_text = False
        for paragraph in list(cell.paragraphs):
            if paragraph.text.strip():
                seen_text = True
                continue
            keep = seen_text and not empty_after_text
            if keep:
                empty_after_text = True
            else:
                paragraph._element.getparent().remove(paragraph._element)
        attorney_paragraph = next(
            (
                paragraph
                for paragraph in cell.paragraphs
                if paragraph.text.strip().startswith("Патентный поверенный")
            ),
            None,
        )
        if attorney_paragraph is not None:
            attorney_paragraph._element.addnext(OxmlElement("w:p"))
        tr_pr = row._tr.get_or_add_trPr()
        height = tr_pr.find(qn("w:trHeight"))
        if height is None:
            height = OxmlElement("w:trHeight")
            tr_pr.append(height)
        height.set(qn("w:val"), "1100")

        # Word keeps a structural paragraph after the final table. Make it
        # minimal so it does not create a visually empty third page.
        table_tail = document.tables[0]._tbl.getnext()
        trailing_paragraph = (
            Paragraph(table_tail, document._body)
            if table_tail is not None and table_tail.tag == qn("w:p")
            else None
        )
        if trailing_paragraph is not None and not trailing_paragraph.text.strip():
            trailing_paragraph.paragraph_format.space_before = Pt(0)
            trailing_paragraph.paragraph_format.space_after = Pt(0)
            trailing_paragraph.paragraph_format.line_spacing = Pt(1)


def _clear_body(document: DocumentType) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_page_break_before(paragraph_element) -> None:
    """Keep every author's two-page block separate from the previous author."""

    paragraph_properties = paragraph_element.find(qn("w:pPr"))
    if paragraph_properties is None:
        paragraph_properties = OxmlElement("w:pPr")
        paragraph_element.insert(0, paragraph_properties)
    if paragraph_properties.find(qn("w:pageBreakBefore")) is None:
        paragraph_properties.append(OxmlElement("w:pageBreakBefore"))


def _strip_comments(path: Path) -> None:
    atomic_write_bytes(path, strip_comments(path.read_bytes()))


def save_consents(
    data: SoftwareConsentData,
    output_path: str | Path,
    template_path: str | Path = DEFAULT_CONSENT_TEMPLATE,
) -> Path:
    if not data.authors:
        raise ValueError("Добавьте хотя бы одного автора.")
    template = Path(template_path)
    if not template.is_file():
        raise FileNotFoundError(f"Не найден шаблон согласия: {template}")

    if len(data.authors) == 1:
        # The normal export path creates one file per author. Editing the
        # template directly avoids loading and cloning the same DOCX twice.
        final = Document(str(template))
        _fill_author_document(final, data, data.authors[0])
    else:
        final = Document(str(template))
        _clear_body(final)
        body = final._element.body
        section_properties = body.sectPr
        for author_index, author in enumerate(data.authors):
            author_document = Document(str(template))
            _fill_author_document(author_document, data, author)
            children = [
                deepcopy(child)
                for child in author_document._element.body
                if child.tag != qn("w:sectPr")
            ]
            if author_index:
                first_paragraph = next(
                    (child for child in children if child.tag == qn("w:p")),
                    None,
                )
                if first_paragraph is not None:
                    _set_page_break_before(first_paragraph)
            for child in children:
                body.insert(body.index(section_properties), child)

    destination = Path(output_path)
    with atomic_output_path(destination) as temporary:
        final.save(str(temporary))
        _strip_comments(temporary)
    return destination


def _safe_author_name(value: str) -> str:
    """Return a file-name-safe FIO on both macOS and Windows."""

    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    return value or "Автор"


def _available_output_path(directory: Path, stem: str, reserved: set[Path]) -> Path:
    candidate = directory / f"{stem}.docx"
    counter = 2
    while candidate.exists() or candidate.is_symlink() or candidate in reserved:
        candidate = directory / f"{stem} ({counter}).docx"
        counter += 1
    reserved.add(candidate)
    return candidate


def save_author_consents(
    data: SoftwareConsentData,
    output_directory: str | Path,
    template_path: str | Path = DEFAULT_CONSENT_TEMPLATE,
) -> list[Path]:
    """Create one two-page consent DOCX per author without overwriting files."""

    if not data.authors:
        raise ValueError("Добавьте хотя бы одного автора.")
    directory = Path(output_directory)
    directory.mkdir(parents=True, exist_ok=True)
    reserved: set[Path] = set()
    destinations = [
        _available_output_path(
            directory,
            f"Согласие {_safe_author_name(author.full_name)}",
            reserved,
        )
        for author in data.authors
    ]

    # Generate the whole batch first, so a failure does not leave a partial set.
    with TemporaryDirectory(prefix=".consents-", dir=directory) as temporary:
        temporary_directory = Path(temporary)
        staged: list[Path] = []
        for index, author in enumerate(data.authors, 1):
            staged_path = temporary_directory / f"{index}.docx"
            save_consents(
                replace(data, authors=[author]),
                staged_path,
                template_path,
            )
            staged.append(staged_path)
        committed: list[Path] = []
        try:
            for source, destination in zip(staged, destinations, strict=True):
                # Recheck immediately before the move. This prevents a file
                # created by another process after reservation from being lost.
                if destination.exists() or destination.is_symlink():
                    raise FileExistsError(
                        f"Файл появился во время сохранения: {destination}"
                    )
                source.replace(destination)
                committed.append(destination)
        except Exception:
            # Every committed path was selected as previously non-existent and
            # was created by this batch, so rolling it back is safe.
            for destination in committed:
                destination.unlink(missing_ok=True)
            raise
    return destinations
