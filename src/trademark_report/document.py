"""Generate client-ready DOCX reports from the retained visual template."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
import sys
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor

from .fees import calculate_fees, class_word, format_rubles
from .models import ConclusionParagraph, ConclusionRun, PERFORMERS, ReportData, SimilarRecord
from .templates import conclusion_paragraphs


def _resource_root() -> Path:
    """Return the project root both from source and from a PyInstaller bundle."""

    bundle_root = getattr(sys, "_MEIPASS", None)
    if getattr(sys, "frozen", False) and bundle_root:
        return Path(bundle_root)
    return Path(__file__).resolve().parents[2]


PROJECT_ROOT = _resource_root()
DEFAULT_TEMPLATE = PROJECT_ROOT / "assets" / "report_template.docx"

BLUE_FILL = "DBE5F1"
GREEN_FILL = "C4D79B"
WARNING_FILL = "FFF200"
GRAY = RGBColor(0x80, 0x80, 0x80)
PINK = "F085AD"


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_font(run, size: float = 11, *, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _format_paragraph(paragraph, *, before=0, after=5, line=1.0, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, *, color="000000", size=4, top=None, bottom=None) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(top if edge == "top" and top else bottom if edge == "bottom" and bottom else size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def _set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(round(Cm(width).emu / 635)))
        grid.append(column)
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)


def _prepare_table(table, widths_cm: list[float]) -> None:
    _set_table_widths(table, widths_cm)
    _set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _clear_cell(cell):
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    return paragraph


def _cell_text(cell, text: str, *, bold=False, italic=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = _clear_cell(cell)
    _format_paragraph(paragraph, after=0, alignment=alignment)
    _set_font(paragraph.add_run(text), size, bold=bold, italic=italic)
    return paragraph


def _add_hyperlink(paragraph, text: str, url: str, *, size=11, bold=False) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attr}"), "Calibri")
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(round(size * 2)))
    properties.append(size_node)
    if bold:
        properties.append(OxmlElement("w:b"))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _add_title(document: Document, report: ReportData) -> None:
    date_paragraph = document.add_paragraph()
    _format_paragraph(date_paragraph, after=16, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    date_run = _set_font(date_paragraph.add_run(f"Дата отчета {report.report_date:%d.%m.%Y}"), 11)
    date_run.font.color.rgb = GRAY

    title = document.add_paragraph()
    _format_paragraph(title, after=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(title.add_run("Отчет об оценке возможности\n"), 11, bold=True)
    _set_font(title.add_run(f"регистрации товарного знака «{report.designation}»"), 11, bold=True)


def _add_summary(document: Document, report: ReportData) -> None:
    table = document.add_table(rows=0, cols=2)
    summary_fields = [("Поисковые запросы", report.search_queries)]
    if report.business_area.strip():
        summary_fields.append(("Сфера деятельности", report.business_area))
    for label, value in summary_fields:
        row = table.add_row()
        _set_cell_shading(row.cells[0], BLUE_FILL)
        _cell_text(row.cells[0], label, bold=True)
        _cell_text(row.cells[1], value or "—")

    row = table.add_row()
    _set_cell_shading(row.cells[0], BLUE_FILL)
    _cell_text(row.cells[0], "Товары и услуги\nсогласно перечню МКТУ", bold=True)
    paragraph = _clear_cell(row.cells[1])
    _format_paragraph(paragraph, after=0)
    if not report.nice_classes:
        _set_font(paragraph.add_run("—"))
    for index, item in enumerate(report.nice_classes):
        if index:
            paragraph = row.cells[1].add_paragraph()
            _format_paragraph(paragraph, after=0)
        _set_font(paragraph.add_run(f"Класс {item.number}: "), 11, bold=True)
        _set_font(paragraph.add_run(item.description), 11)
    _prepare_table(table, [5.3, 11.7])

    spacer = document.add_paragraph()
    _format_paragraph(spacer, after=3)

    grounds = document.add_table(rows=2, cols=2)
    _set_cell_shading(grounds.cell(0, 0), BLUE_FILL)
    _set_cell_shading(grounds.cell(1, 0), BLUE_FILL)
    _cell_text(grounds.cell(0, 0), "Абсолютные основания\nдля отказа в регистрации обозначения", bold=True)
    absolute_text = (
        report.absolute_grounds_text.strip() or "Обозначение может быть признано не соответствующим требованиям законодательства."
        if report.has_absolute_grounds
        else "Отсутствуют"
    )
    _cell_text(grounds.cell(0, 1), absolute_text)
    _cell_text(grounds.cell(1, 0), "Относительные основания\nдля отказа в регистрации обозначения", bold=True)

    relative_cell = grounds.cell(1, 1)
    relative_paragraph = _clear_cell(relative_cell)
    _format_paragraph(relative_paragraph, after=5)
    selected = [item for item in report.relative_options if item != "Отсутствуют"]
    relative_text = "; ".join(selected) if selected else "Отсутствуют"
    _set_font(relative_paragraph.add_run(relative_text))
    date_one = relative_cell.add_paragraph()
    _format_paragraph(date_one, before=8, after=1)
    run = _set_font(date_one.add_run(f"База товарных знаков обновлена {report.trademarks_database_date}"), 11, italic=True)
    run.font.color.rgb = GRAY
    date_two = relative_cell.add_paragraph()
    _format_paragraph(date_two, after=0)
    run = _set_font(date_two.add_run(f"База заявок на товарные знаки обновлена {report.applications_database_date}"), 11, italic=True)
    run.font.color.rgb = GRAY
    _prepare_table(grounds, [5.3, 11.7])


def _add_conclusion(document: Document, report: ReportData) -> None:
    header = document.add_table(rows=1, cols=1)
    _set_cell_shading(header.cell(0, 0), GREEN_FILL)
    _cell_text(header.cell(0, 0), "ЗАКЛЮЧЕНИЕ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _prepare_table(header, [17.0])

    box = document.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    content = report.conclusion_content
    if content is None:
        content = []
        for role, text in conclusion_paragraphs(report):
            content.append(
                ConclusionParagraph(
                    runs=[
                        ConclusionRun(
                            text=text,
                            bold=role in {"bold", "warning"},
                            italic=role == "italic",
                            highlighted=role == "warning",
                        )
                    ],
                    list_item=role == "list",
                )
            )

    first = True
    list_index = 0
    for item in content:
        paragraph = _clear_cell(cell) if first else cell.add_paragraph()
        first = False
        _format_paragraph(paragraph, before=2, after=7, line=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if item.list_item:
            marker = chr(ord("a") + list_index)
            list_index += 1
            _set_font(paragraph.add_run(f"{marker}) "), 11)
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        for content_run in item.runs:
            run = _set_font(
                paragraph.add_run(content_run.text),
                11,
                bold=content_run.bold,
                italic=content_run.italic,
            )
            if content_run.highlighted:
                highlight = OxmlElement("w:highlight")
                highlight.set(qn("w:val"), "yellow")
                run._r.get_or_add_rPr().append(highlight)
    _prepare_table(box, [17.0])

    probability = document.add_table(rows=max(len(report.probabilities), 1), cols=2)
    entries = report.probabilities or []
    if not entries:
        from .models import ProbabilityEntry

        entries = [ProbabilityEntry("—")]
    for index, entry in enumerate(entries):
        label = "Вероятность регистрации\nобозначения"
        if entry.subject.strip():
            label += f" {entry.subject.strip()}"
        _set_cell_shading(probability.cell(index, 0), GREEN_FILL)
        _cell_text(probability.cell(index, 0), label, bold=True)
        _cell_text(probability.cell(index, 1), entry.value, bold=True)
    _prepare_table(probability, [5.3, 11.7])

    signature = document.add_paragraph()
    _format_paragraph(signature, before=8, after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    performer_name = PERFORMERS.get(report.performer, report.performer)
    _set_font(signature.add_run(f"С уважением,\n{performer_name}\n8-800-222-90-53\n"), 11)
    _add_hyperlink(signature, "www.patentural.ru", "https://www.patentural.ru/", size=11)


def _add_section_caption(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_cell_shading(table.cell(0, 0), BLUE_FILL)
    _cell_text(table.cell(0, 0), text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _prepare_table(table, [17.0])


def _add_record_table(document: Document, record: SimilarRecord) -> None:
    if record.kind == "international":
        fields = [
            ("Номер регистрации:", record.number),
            ("Дата регистрации:", record.relevant_date),
            ("Правообладатель:", record.owner_or_applicant),
            ("Однородные классы МКТУ:", record.related_classes),
        ]
    elif record.kind == "application":
        fields = [
            ("Номер заявки:", record.number),
            ("Статус:", record.status),
            ("Дата подачи:", record.relevant_date),
            ("Заявитель:", record.owner_or_applicant),
            ("Однородные классы МКТУ:", record.related_classes),
        ]
    else:
        fields = [
            ("Номер регистрации:", record.number),
            ("Дата приоритета:", record.relevant_date),
            ("Правообладатель:", record.owner_or_applicant),
        ]
        if record.unprotected_element.strip():
            fields.append(("Неохраняемый элемент:", record.unprotected_element))
        fields.append(("Однородные классы МКТУ:", record.related_classes))

    table = document.add_table(rows=len(fields), cols=3)
    image_cell = table.cell(0, 0).merge(table.cell(len(fields) - 1, 0))
    image_paragraph = _clear_cell(image_cell)
    _format_paragraph(image_paragraph, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if record.image_bytes:
        try:
            image_paragraph.add_run().add_picture(BytesIO(record.image_bytes), width=Cm(4.2))
        except Exception:
            _set_font(image_paragraph.add_run("Изображение не удалось вставить"), 11, italic=True)
    elif record.display_name:
        _set_font(image_paragraph.add_run(record.display_name), 11, bold=True)

    for index, (label, value) in enumerate(fields):
        _cell_text(table.cell(index, 1), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        value_paragraph = _clear_cell(table.cell(index, 2))
        _format_paragraph(value_paragraph, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if index == 0 and record.source_url.strip() and record.number.strip():
            _add_hyperlink(value_paragraph, record.number, record.source_url, size=11, bold=False)
        else:
            _set_font(value_paragraph.add_run(value or "—"), 11)
    _prepare_table(table, [5.4, 5.2, 6.4])
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, after=1)


def _add_appendix(document: Document, report: ReportData) -> None:
    heading = document.add_paragraph()
    _format_paragraph(heading, after=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_font(heading.add_run("Приложение 1"), 11, bold=True)

    sections = (
        ("МЕЖДУНАРОДНЫЕ ТОВАРНЫЕ ЗНАКИ, ЗАРЕГИСТРИРОВАННЫЕ В РФ", report.international_marks),
        ("ТОВАРНЫЕ ЗНАКИ РФ", report.russian_marks),
        ("ЗАЯВКИ НА ТОВАРНЫЕ ЗНАКИ", report.applications),
    )
    for title, records in sections:
        if not records:
            continue
        _add_section_caption(document, title)
        for record in records:
            _add_record_table(document, record)


def _add_fees(document: Document, report: ReportData) -> None:
    class_count = max(len(report.nice_classes), 1)
    fees = calculate_fees(class_count)

    title = document.add_paragraph()
    _format_paragraph(title, after=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(title.add_run("ГОСУДАРСТВЕННЫЕ ПОШЛИНЫ"), 11, bold=True)

    table = document.add_table(rows=4, cols=4)
    headers = ("№", "Пошлины", "Сумма, руб.", "Сроки оплаты")
    for index, value in enumerate(headers):
        _set_cell_shading(table.cell(0, index), BLUE_FILL)
        _cell_text(table.cell(0, index), value, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_text(table.cell(1, 0), "1", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(
        table.cell(1, 1),
        f"Государственная пошлина за подачу заявки и проведение экспертизы: "
        f"{class_count} {class_word(class_count)} МКТУ",
    )
    _cell_text(
        table.cell(1, 2),
        f"{format_rubles(fees.filing)} + 500 за каждый товар/услугу свыше 10",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _cell_text(table.cell(1, 3), "В течение 2-х месяцев с даты подачи заявки", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_text(table.cell(2, 0), "2", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(
        table.cell(2, 1),
        f"Государственная пошлина за регистрацию товарного знака и получение электронного "
        f"свидетельства: {class_count} {class_word(class_count)} МКТУ",
    )
    _cell_text(table.cell(2, 2), format_rubles(fees.registration), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(table.cell(2, 3), "В течение 2-х месяцев с даты решения о регистрации", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    merged = table.cell(3, 0).merge(table.cell(3, 1))
    _cell_text(merged, "ИТОГО", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    total_cell = table.cell(3, 2).merge(table.cell(3, 3))
    _cell_text(
        total_cell,
        format_rubles(fees.total),
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _prepare_table(table, [0.8, 7.4, 4.0, 4.8])

    note = document.add_paragraph()
    _format_paragraph(note, before=12, after=5)
    _set_font(note.add_run("Примечание:"), 11)
    notes = (
        "Срок действия товарного знака — 10 лет с даты подачи заявки.",
        "Пошлина за подачу заявки может быть оплачена в течение двух месяцев с даты подачи "
        "заявки. Однако экспертиза начнется только после оплаты государственной пошлины.",
        "Существует возможность получения свидетельства о регистрации товарного знака в "
        "бумажном виде. Дополнительная пошлина за бумажное свидетельство составляет 3 000 руб.",
    )
    for index, text in enumerate(notes, 1):
        paragraph = document.add_paragraph()
        _format_paragraph(paragraph, after=3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        _set_font(paragraph.add_run(f"{index}. {text}"), 11)

    for _ in range(2):
        _format_paragraph(document.add_paragraph(), after=0)

    social = document.add_table(rows=1, cols=1)
    _set_table_borders(social, color=PINK, size=0, top=24, bottom=24)
    cell = social.cell(0, 0)
    paragraph = _clear_cell(cell)
    _format_paragraph(paragraph, before=4, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(paragraph.add_run("Подписывайтесь на нас в соцсетях!"), 11, bold=True, italic=True)
    paragraph = cell.add_paragraph()
    _format_paragraph(paragraph, after=5)
    _set_font(
        paragraph.add_run(
            "Рассказываем об интеллектуальной собственности просто и с юмором. Активно делимся "
            "яркими моментами, полезной информацией и эксклюзивными новостями."
        ),
        11,
        italic=True,
    )
    links = (
        ("Telegram", "https://t.me/patentural"),
        ("YouTube", "https://youtube.com/channel/UCSocRtxH36x9FeukIXrAe7A"),
        ("Rutube", "https://rutube.ru/u/patentural/"),
        ("ВКонтакте", "https://vk.com/privilege_cp"),
        ("Instagram", "https://www.instagram.com/patentural/"),
    )
    for text, url in links:
        paragraph = cell.add_paragraph()
        _format_paragraph(paragraph, after=0)
        paragraph.paragraph_format.left_indent = Cm(1.8)
        _add_hyperlink(paragraph, text, url, size=11)
    paragraph = cell.add_paragraph()
    _format_paragraph(paragraph, before=3, after=4)
    _set_font(
        paragraph.add_run(
            "*Meta (Instagram) признана экстремистской организацией и запрещена на территории РФ*"
        ),
        11,
        italic=True,
    )
    _prepare_table(social, [17.0])
    _set_table_borders(social, color=PINK, size=0, top=24, bottom=24)


def _strip_comment_parts(docx_bytes: bytes) -> bytes:
    output = BytesIO()
    removable = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsExtensible.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    }
    with ZipFile(BytesIO(docx_bytes), "r") as source, ZipFile(output, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename in removable:
                continue
            data = source.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                text = re.sub(
                    r'<Relationship\b[^>]+Type="[^"]+/(?:comments|commentsExtended|commentsExtensible|commentsIds|people)"[^>]*/>',
                    "",
                    text,
                )
                data = text.encode("utf-8")
            elif item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                text = re.sub(r'<Override\b[^>]+PartName="/word/(?:comments[^"/]*|people)\.xml"[^>]*/>', "", text)
                data = text.encode("utf-8")
            target.writestr(item, data)
    return output.getvalue()


def generate_report(report: ReportData, template_path: str | Path = DEFAULT_TEMPLATE) -> bytes:
    """Create a DOCX report and return its bytes."""

    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Не найден Word-шаблон: {template_path}")
    if not report.designation.strip() or not report.search_queries.strip():
        raise ValueError("Обозначение и поисковые запросы обязательны.")

    document = Document(template_path)
    _clear_body(document)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(5)

    _add_title(document, report)
    _add_summary(document, report)
    _add_page_break(document)
    _add_conclusion(document, report)
    if report.has_appendix:
        _add_page_break(document)
        _add_appendix(document, report)
    _add_page_break(document)
    _add_fees(document, report)

    buffer = BytesIO()
    document.save(buffer)
    return _strip_comment_parts(buffer.getvalue())


def save_report(report: ReportData, output_path: str | Path, template_path: str | Path = DEFAULT_TEMPLATE) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(generate_report(report, template_path))
    return output_path
