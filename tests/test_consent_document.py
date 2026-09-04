from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from trademark_report.consent_document import (
    DEFAULT_CONSENT_TEMPLATE,
    save_author_consents,
    save_consents,
)
from trademark_report.software_models import (
    APPLICANT_INDIVIDUAL,
    SoftwareAuthor,
    SoftwareConsentData,
)


def _data() -> SoftwareConsentData:
    return SoftwareConsentData(
        program_name="Тестовая программа",
        applicant_name="ООО «Тест»",
        applicant_address="Россия, г. Москва, ул. Тестовая, д. 1",
        inn="7707083893",
        ogrn="1027700132195",
        application_number="2026888888",
        document_date=date(2026, 8, 26),
        authors=[
            SoftwareAuthor(
                full_name="Иванов Иван Иванович",
                birth_date="01.01.1980",
                address="Россия, г. Москва, ул. Первая, д. 1",
                passport_series="4510",
                passport_number="123456",
                passport_issue_date="02.02.2020",
                passport_issuer=(
                    "ГУ МВД России по г. Москве\n"
                    "Код подразделения: 770-068\nДата выдачи:"
                ),
                creative_contribution="Разработка алгоритма",
            ),
            SoftwareAuthor(
                full_name="Петров Петр Петрович",
                birth_date="03.03.1981",
                address="Россия, г. Москва, ул. Вторая, д. 2",
                passport_series="4511",
                passport_number="654321",
                passport_issue_date="04.04.2021",
                passport_issuer="ГУ МВД России по г. Москве",
                creative_contribution="Написание исходного текста программы",
            ),
        ],
    )


def test_combined_consents_preserve_template_and_remove_comments(tmp_path):
    output = tmp_path / "consents.docx"
    save_consents(_data(), output)

    document = Document(output)
    assert len(document.tables) == 2
    assert len(document.comments) == 0
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Иванов Иван Иванович" in text
    assert "Петров Петр Петрович" in text
    assert "Тестовая программа" in text
    assert "1027700132195" in text
    assert "26.08.2026" in text
    assert "2026888888" not in text
    assert "№ заявки" not in text
    assert "Заявка №" not in text
    assert "регистрационного номера заявки" not in text
    assert "серия «4510» номер «123456» выдан 02.02.2020" in text
    assert "выдан «02.02.2020" not in text
    assert "Код подразделения" not in text
    assert "Дата выдачи:" not in text
    assert "«Разработка алгоритма»" not in text
    assert "Разработка алгоритма" in text
    assert "ФИО" not in text
    recipient_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "В Федеральную службу"
    ]
    assert len(recipient_paragraphs) == 2
    assert recipient_paragraphs[0].paragraph_format.page_break_before is not True
    assert recipient_paragraphs[1].paragraph_format.page_break_before is True

    identity_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Документ, удостоверяющий личность")
    ]
    assert identity_paragraphs
    assert all("\n" not in paragraph.text for paragraph in identity_paragraphs)
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        for paragraph in identity_paragraphs
    )
    table_widths = [
        int(table._tbl.tblPr.find(qn("w:tblW")).get(qn("w:w")))
        for table in document.tables
    ]
    assert table_widths == [10244, 10244]

    author_signature_cell = document.tables[0].rows[4].cells[0]
    author_signature_index = next(
        index
        for index, paragraph in enumerate(author_signature_cell.paragraphs)
        if paragraph.text.startswith("Подпись автора:")
    )
    assert sum(
        not paragraph.text.strip()
        for paragraph in author_signature_cell.paragraphs[:author_signature_index]
    ) == 3

    attorney_cell = document.tables[0].rows[5].cells[0]
    assert attorney_cell.paragraphs[0].text.startswith("Патентный поверенный")
    assert not attorney_cell.paragraphs[1].text.strip()
    assert not attorney_cell.paragraphs[2].text.strip()
    assert attorney_cell.paragraphs[3].text == "26.08.2026"
    title_row_height = document.tables[0].rows[2]._tr.trPr.find(qn("w:trHeight"))
    assert title_row_height.get(qn("w:val")) == "900"

    table_element = document.tables[0]._tbl
    page_break_paragraph = table_element.getprevious()
    assert page_break_paragraph.tag == qn("w:p")
    assert any(
        item.get(qn("w:type")) == "page"
        for item in page_break_paragraph.findall(f".//{qn('w:br')}")
    )

    for table in document.tables:
        following_element = table._tbl.getnext()
        assert following_element.tag == qn("w:p")
        spacing = following_element.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
        assert spacing.get(qn("w:before")) == "0"
        assert spacing.get(qn("w:after")) == "0"
        assert spacing.get(qn("w:line")) == "20"

    with ZipFile(output) as archive:
        names = set(archive.namelist())
        assert "word/comments.xml" not in names
        assert not any("comments" in name.casefold() for name in names)
        assert b"commentReference" not in archive.read("word/document.xml")


def test_consent_template_is_bundled():
    assert DEFAULT_CONSENT_TEMPLATE.is_file()


def test_anonymous_author_choice_is_checked_in_generated_document(tmp_path):
    output = tmp_path / "anonymous-consent.docx"
    data = _data()
    data.authors_will_be_mentioned = False

    save_consents(data, output)

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "☐ упоминать его под своим именем     ☒ не упоминать его (анонимно)" in text
    assert "☒ упоминать его под своим именем     ☐ не упоминать его (анонимно)" not in text


def test_separate_consent_is_created_for_each_author(tmp_path):
    paths = save_author_consents(_data(), tmp_path)

    assert [path.name for path in paths] == [
        "Согласие Иванов Иван Иванович.docx",
        "Согласие Петров Петр Петрович.docx",
    ]
    first = Document(str(paths[0]))
    second = Document(str(paths[1]))
    first_text = "\n".join(paragraph.text for paragraph in first.paragraphs)
    second_text = "\n".join(paragraph.text for paragraph in second.paragraphs)
    assert len(first.tables) == 1
    assert len(second.tables) == 1
    assert "Иванов Иван Иванович" in first_text
    assert "Петров Петр Петрович" not in first_text
    assert "Петров Петр Петрович" in second_text
    assert "Иванов Иван Иванович" not in second_text


def test_separate_consents_do_not_overwrite_existing_files(tmp_path):
    existing = tmp_path / "Согласие Иванов Иван Иванович.docx"
    existing.write_bytes(b"existing")

    paths = save_author_consents(_data(), tmp_path)

    assert existing.read_bytes() == b"existing"
    assert paths[0].name == "Согласие Иванов Иван Иванович (2).docx"


def test_separate_consent_batch_is_rolled_back_if_commit_fails(monkeypatch, tmp_path):
    original_replace = Path.replace

    def failing_replace(source, destination):
        if Path(destination).name == "Согласие Петров Петр Петрович.docx":
            raise OSError("simulated disk failure")
        return original_replace(source, destination)

    monkeypatch.setattr(Path, "replace", failing_replace)

    try:
        save_author_consents(_data(), tmp_path)
    except OSError:
        pass
    else:
        raise AssertionError("Expected a simulated save failure")

    assert not list(tmp_path.glob("Согласие *.docx"))


def test_individual_applicant_document_uses_inn_without_ogrn(tmp_path):
    output = tmp_path / "individual.docx"
    data = _data()
    data.applicant_type = APPLICANT_INDIVIDUAL
    data.applicant_name = "Анфиногенов Семен Васильевич"
    data.applicant_address = "Россия, г. Екатеринбург, ул. Рощинская, д. 59"
    data.inn = "667803703833"
    data.ogrn = ""

    save_consents(data, output)

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Анфиногенов Семен Васильевич" in text
    assert "ИНН: 667803703833" in text
    assert "ОГРН:" not in text
    assert "125993" in text
    assert "119991" not in text
