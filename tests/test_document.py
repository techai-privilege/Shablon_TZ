from datetime import date
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from trademark_report.document import _resource_root, generate_report
from trademark_report.models import (
    ConclusionParagraph,
    ConclusionRun,
    ProbabilityEntry,
    ReportData,
    ReportNiceClass,
    SimilarRecord,
)


def test_generate_report_contains_expected_content_and_no_comments():
    report = ReportData(
        designation="FELICHE",
        search_queries="FELICHE, ФЕЛИЧЕ",
        business_area="Розничная торговля",
        report_date=date(2026, 8, 24),
        nice_classes=[ReportNiceClass("35", "услуги розничной торговли")],
        trademarks_database_date="24.08.2026",
        applications_database_date="24.08.2026",
        performer="Маша",
        probabilities=[ProbabilityEntry(value="ВЫШЕ СРЕДНЕЙ")],
        russian_marks=[
            SimilarRecord(
                kind="russian",
                source_url=(
                    "https://new.fips.ru/registers-doc-view/fips_servlet?"
                    "DB=RUTM&DocNumber=1227836&TypeFile=html"
                ),
                display_name="IGS",
                number="1227836",
                relevant_date="03.02.2026",
                owner_or_applicant="ООО «Тест»",
                related_classes="35",
            )
        ],
    )

    generated = generate_report(report)
    assert generated.startswith(b"PK")

    document = Document(BytesIO(generated))
    body_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "FELICHE" in body_text
    assert "1227836" in body_text
    assert "Дукки Алина" not in body_text
    assert "Шмыкова Мария" in body_text
    assert "ГОСУДАРСТВЕННЫЕ ПОШЛИНЫ" in body_text
    assert "17 000 + 500 за каждый товар/услугу свыше 10" in body_text
    assert "18 000" in body_text

    with ZipFile(BytesIO(generated)) as archive:
        names = set(archive.namelist())
        assert "word/comments.xml" not in names
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "DocNumber=1227836" in relationships


def test_resource_root_uses_pyinstaller_bundle(monkeypatch, tmp_path):
    monkeypatch.setattr("trademark_report.document.sys.frozen", True, raising=False)
    monkeypatch.setattr("trademark_report.document.sys._MEIPASS", str(tmp_path), raising=False)

    assert _resource_root() == tmp_path


def test_fees_table_uses_class_count_reference_and_no_manual_excess_total():
    report = ReportData(
        designation="TEST",
        search_queries="TEST",
        nice_classes=[ReportNiceClass(str(index).zfill(2), "описание") for index in range(1, 11)],
        trademarks_database_date="24.08.2026",
        applications_database_date="24.08.2026",
    )

    document = Document(BytesIO(generate_report(report)))
    fees_table = next(
        table
        for table in document.tables
        if len(table.columns) >= 2 and table.cell(0, 1).text.strip() == "Пошлины"
    )
    assert fees_table.cell(1, 2).text.strip() == (
        "48 500 + 500 за каждый товар/услугу свыше 10"
    )
    assert fees_table.cell(2, 2).text.strip() == "28 000"
    assert fees_table.cell(3, 2).text.strip() == "…"


def test_edited_conclusion_content_is_used_in_docx_with_formatting():
    report = ReportData(
        designation="TEST",
        search_queries="TEST",
        nice_classes=[ReportNiceClass("35", "услуги")],
        trademarks_database_date="24.08.2026",
        applications_database_date="24.08.2026",
        conclusion_content=[
            ConclusionParagraph(
                runs=[ConclusionRun("Исправленный вывод эксперта", bold=True)]
            ),
            ConclusionParagraph(
                runs=[ConclusionRun("Особое предупреждение", highlighted=True)]
            ),
        ],
    )

    document = Document(BytesIO(generate_report(report)))
    conclusion_box = next(
        table for table in document.tables if "Исправленный вывод эксперта" in table.cell(0, 0).text
    )
    paragraphs = conclusion_box.cell(0, 0).paragraphs
    assert paragraphs[0].runs[0].bold is True
    assert "Особое предупреждение" in paragraphs[1].text
    assert paragraphs[1].runs[0]._r.xpath("./w:rPr/w:highlight")
