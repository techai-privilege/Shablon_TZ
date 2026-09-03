from pathlib import Path

from docx import Document

from trademark_report.questionnaire import parse_questionnaire


def _questionnaire(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=0, cols=2)
    values = [
        ("РАЗДЕЛ I – СВЕДЕНИЯ О ПО ДЛЯ РОСПАТЕНТА", ""),
        ("Наименование программы", "Тестовая программа"),
        ("Наименование заявителя и ИНН", "ООО «Тест», ИНН 7707083893"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные авторов",
            "Автор 1:\n1. Иванов Иван Иванович\n2. 01.01.1980\n"
            "3. Россия, г. Москва, ул. Тестовая, д. 1\n4. 4510 123456\n"
            "5. ГУ МВД России по г. Москве 02.02.2020\n\n"
            "Автор 2:\n1. Петров Петр Петрович\n2. 03.03.1981\n"
            "3. Россия, г. Москва, ул. Вторая, д. 2\n4. 4511 654321\n"
            "5. ГУ МВД России по г. Москве 04.04.2021",
        ),
        (
            "Творческий вклад авторов",
            "Автор 1:\n- Разработка алгоритма\nАвтор 2:\n- Написание исходного текста программы",
        ),
    ]
    for label, value in values:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    ignored = document.add_table(rows=2, cols=2)
    ignored.cell(0, 0).text = "РАЗДЕЛ II – СВЕДЕНИЯ О ПО ДЛЯ РЕЕСТРА"
    ignored.cell(1, 0).text = "Наименование программы"
    ignored.cell(1, 1).text = "Это значение нельзя импортировать"
    document.save(path)


def test_parser_uses_only_section_one_and_splits_authors(tmp_path):
    path = tmp_path / "questionnaire.docx"
    _questionnaire(path)

    result = parse_questionnaire(path)

    assert result.data.program_name == "Тестовая программа"
    assert result.data.applicant_name == "ООО «Тест»"
    assert result.data.inn == "7707083893"
    assert result.data.declared_author_count == 2
    assert len(result.data.authors) == 2
    assert result.data.authors[0].full_name == "Иванов Иван Иванович"
    assert result.data.authors[0].passport_series == "4510"
    assert result.data.authors[0].passport_number == "123456"
    assert result.data.authors[0].creative_contribution == "Разработка алгоритма"
    assert result.data.authors[1].full_name == "Петров Петр Петрович"
    assert result.data.authors[1].birth_date == "03.03.1981"
    assert "Это значение нельзя импортировать" not in result.data.program_name


def test_parser_removes_service_fields_from_passport_issuer(tmp_path):
    path = tmp_path / "passport-service-fields.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("РАЗДЕЛ I – СВЕДЕНИЯ О ПО ДЛЯ РОСПАТЕНТА", ""),
        ("Наименование программы", "Тестовая программа"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные авторов",
            "Автор 1:\n1. Иванов Иван Иванович\n2. 01.01.1980\n"
            "3. Россия, г. Москва\n4. 4510 123456\n"
            "5. ГУ МВД России по г. Москве 02.02.2020 "
            "Код подразделения: 770-068 Дата выдачи:",
        ),
        ("Творческий вклад авторов", "Автор 1:\n- Разработка алгоритма"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.passport_issue_date == "02.02.2020"
    assert author.passport_issuer == "ГУ МВД России по г. Москве"


def test_parser_reads_section_one_split_across_two_tables(tmp_path):
    path = tmp_path / "split-section.docx"
    document = Document()
    first = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "1YES Analitik"),
        ("Область (сфера) применения", "Аналитика"),
    ):
        cells = first.add_row().cells
        cells[0].text = label
        cells[1].text = value
    second = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование заявителя и ИНН", "ООО «Старшип.про», ИНН 1683025264"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные автора 1",
            "- УРБАН АЛЕКСАНДР СЕРГЕЕВИЧ\n- 25.03.1988\n"
            "- 141407, ГОРОД ХИМКИ, ПРОСПЕКТ ЮБИЛЕЙНЫЙ, ДОМ 49\n"
            "- 4615 931690, ТП №1 ОУФМС РОССИИ ПО МОСКОВСКОЙ ОБЛАСТИ, 25.06. 2015",
        ),
        ("Описание творческого вклада автора 1", "- разработка спецификаций программы"),
    ):
        cells = second.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)
    author = result.data.authors[0]

    assert result.data.applicant_name == "ООО «Старшип.про»"
    assert result.data.inn == "1683025264"
    assert author.full_name == "Урбан Александр Сергеевич"
    assert author.passport_series == "4615"
    assert author.passport_number == "931690"
    assert author.passport_issue_date == "25.06.2015"
    assert author.passport_issuer.startswith("ТП №1 ОУФМС")


def test_parser_reads_labeled_passport_line_and_address(tmp_path):
    path = tmp_path / "labeled-passport.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Лина"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада автора 1",
            "Автор: Баскаков Александр Сергеевич\n"
            "- дата рождения: 15.05.2002\n"
            "- адрес места жительства: г. Москва, ул. Черкизовская Б., д. 32\n"
            "- серия и номер паспорта, дата выдачи и выдавший орган: "
            "4522 868797 02.06.2022 ГУ МВД России по г. Москве",
        ),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада автора 1",
            "- разработка всей программы в целом",
        ),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.address == "г. Москва, ул. Черкизовская Б., д. 32"
    assert author.passport_series == "4522"
    assert author.passport_number == "868797"
    assert author.passport_issue_date == "02.06.2022"
    assert author.passport_issuer == "ГУ МВД России по г. Москве"


def test_parser_supports_foreign_passport_and_textual_issue_date(tmp_path):
    path = tmp_path / "passport-variants.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    rows = (
        ("Наименование программы", "КИСС"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные Автора 1",
            "Скориков Максим Сергеевич\n15.07.1991\n"
            "420500, Республика Татарстан, г. Иннополис\n"
            "Паспорт РК № 10930256\nвыдан МИД РК, 29.03.2017",
        ),
        ("Описание творческого вклада автора 1", "- разработка программы в целом"),
        (
            "Паспортные данные Автора 2",
            "Корнеев Кирилл Андреевич\n03.09.1997 года рождения\n"
            "664009, Иркутская область, город Иркутск\nПаспорт: 2517 461759\n"
            "Отделом УФМС России по Иркутской области, выдан «14» ноября 2017 г.",
        ),
        ("Описание творческого вклада автора 2", "- разработка программы в целом"),
    )
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    authors = parse_questionnaire(path).data.authors

    assert (authors[0].passport_series, authors[0].passport_number) == ("РК", "10930256")
    assert authors[0].address.startswith("420500")
    assert authors[1].passport_issue_date == "14.11.2017"
    assert authors[1].passport_issuer.startswith("Отделом УФМС")


def test_parser_separates_contribution_from_passport_and_address(tmp_path):
    path = tmp_path / "combined-passport-contribution.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Пересвет"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада автора 1",
            "Кобелев Денис Александрович, 11.12.1986 г.р., паспорт РФ серии "
            "5707 №013139, выдан 16.01.2007 г. Управлением внутренних дел города Кунгура, "
            "зарегистрированного по адресу: Пермский край, г. Кунгур, ул. Новая, д. 16.\n"
            "- Разработка системной архитектуры\n- Разработка интерфейсов",
        ),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.address == "Пермский край, г. Кунгур, ул. Новая, д. 16."
    assert author.passport_issuer == "Управлением внутренних дел города Кунгура"
    assert author.creative_contribution == "Разработка системной архитектуры; Разработка интерфейсов"


def test_parser_reads_program_name_above_section_one_table(tmp_path):
    path = tmp_path / "extended-questionnaire.docx"
    document = Document()
    document.add_paragraph("ПО «Складские системы RMS»")
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("РАЗДЕЛ I – СВЕДЕНИЯ О ПО ДЛЯ РОСПАТЕНТА", ""),
        ("Область (сфера) применения ПО", "Складская логистика"),
        ("ИНН заявителя", "6321416298"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert result.data.program_name == "Складские системы RMS"


def test_parser_keeps_full_applicant_name_from_questionnaire(tmp_path):
    path = tmp_path / "full-applicant-name.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Программа управления оборудованием"),
        ("Программа создана за счет средств:", "Собственных"),
        (
            "Наименование заявителя и ИНН",
            "Общество с ограниченной ответственностью Производственное "
            "объединение «Комплекс»,\nИНН: 6658460826",
        ),
        ("Количество авторов", "0"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert result.data.applicant_name == (
        "Общество с ограниченной ответственностью Производственное "
        "объединение «Комплекс»"
    )
    assert result.data.questionnaire_profile_id == "cp_software_2026"


def test_parser_joins_wrapped_author_address_until_passport(tmp_path):
    path = tmp_path / "wrapped-author-address.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "СлотПлан"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1",
            "Сумароков Константин Иванович, 15.10.1985\n"
            "РОССИЯ, 454092, Челябинская обл., г. Челябинск,\n"
            "ул. Воровского, д. 21, кв. 39\n"
            "Паспорт: 75 05 865736 Выдан: Управление внутренних дел "
            "по Центральному району города Челябинска 18.04.2006",
        ),
        ("Описание творческого вклада Автора 1", "Разработка архитектуры"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert result.data.authors[0].address == (
        "РОССИЯ, 454092, Челябинская обл., г. Челябинск, "
        "ул. Воровского, д. 21, кв. 39"
    )


def test_parser_supports_issuer_in_separate_issued_by_authority_line(tmp_path):
    path = tmp_path / "issued-by-authority.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "ATLAS MONITORING"),
        ("Программа создана за счет средств:", "Собственных"),
        (
            "Наименование заявителя и ИНН",
            "ООО «Электронные и программные системы», ИНН 6658302724",
        ),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1",
            "ФИО - Пахомов Артем Валерьевич\n"
            "Дата рождения - 10.02.2004\n"
            "Адрес места жительства - г. Екатеринбург, ул. Академика Бардина, д. 41\n"
            "Серия и номер паспорта – 65 24 007573\n"
            "дата выдачи – 02.04.2024\n"
            "выдавший орган – ГУ МВД РОССИИ ПО СВЕРДЛОВСКОЙ ОБЛАСТИ",
        ),
        ("Описание творческого вклада Автора 1", "Разработка программы"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    author = result.data.authors[0]
    assert author.passport_issue_date == "02.04.2024"
    assert author.passport_issuer == "ГУ МВД РОССИИ ПО СВЕРДЛОВСКОЙ ОБЛАСТИ"
    assert not any("кем выдан" in warning for warning in result.warnings)


def test_parser_removes_issuer_field_labels_from_value(tmp_path):
    path = tmp_path / "issuer-field-label.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Пример"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1",
            "Иванов Иван Иванович, 01.01.1990\n"
            "серия 6501, номер 123456, дата выдачи 01.02.2010, "
            "орган, выдавший документ: орган МВД РОССИИ\n"
            "адрес: г. Москва, д. 1",
        ),
        ("Описание творческого вклада Автора 1", "Разработка программы"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.passport_issuer == "МВД РОССИИ"


def test_parser_extracts_issuer_before_date_and_ignores_subdivision_code(tmp_path):
    path = tmp_path / "issuer-before-date.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "LLM Agent"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные Автора 1",
            "Горбатюк Дмитрий Петрович\n14.02.1977\n"
            "Г. Екатеринбург ул. Белинского д.161 кв 171,\n"
            "паспорт РФ 65 21 463031 ГУ МВД России по Свердловской области "
            "22.02.2022 660-003",
        ),
        ("Описание творческого вклада Автора 1", "Разработка алгоритма"),
        (
            "Паспортные данные Автора 2",
            "Горбатюк Олег Дмитриевич\n20.06.2004\n"
            "Г. Екатеринбург ул. Белинского д.161 кв 171\n"
            "Паспорт 65 24 080095 ГУ МВД России по Свердловской области "
            "06.08.2024 660-002",
        ),
        ("Описание творческого вклада Автора 2", "Разработка программы"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert [author.passport_issuer for author in result.data.authors] == [
        "ГУ МВД России по Свердловской области",
        "ГУ МВД России по Свердловской области",
    ]
    assert [author.passport_issue_date for author in result.data.authors] == [
        "22.02.2022",
        "06.08.2024",
    ]


def test_parser_splits_grouped_authors_with_their_contributions(tmp_path):
    path = tmp_path / "grouped-authors.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Программный комплекс"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные Автора 1\nОписание творческого вклада Автора 1",
            "1. Иванов Иван Иванович\n"
            "01.01.1990 г.р., паспорт 6501 123456, выдан МВД России "
            "01.02.2010, зарегистрирован г. Екатеринбург, ул. Ленина, д. 1\n"
            "- Разработка архитектуры программы\n"
            "2. Петров Петр Петрович\n"
            "02.02.1992 г.р., паспорт выдан УФМС России 03.03.2012, "
            "Код подразделения 660-001, серия 6502 No 654321, "
            "адрес: г. Екатеринбург, ул. Мира, д. 2\n"
            "- Написание исходного кода программы",
        ),
        (
            "Паспортные данные Автора 1\nОписание творческого вклада Автора 1",
            "Иванов Иван Иванович - Разработка архитектуры программы.\n"
            "Петров Петр Петрович - Написание исходного кода программы.",
        ),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert len(result.data.authors) == 2
    first, second = result.data.authors
    assert first.full_name == "Иванов Иван Иванович"
    assert first.creative_contribution == "Разработка архитектуры программы"
    assert first.address == "г. Екатеринбург, ул. Ленина, д. 1"
    assert second.full_name == "Петров Петр Петрович"
    assert second.creative_contribution == "Написание исходного кода программы"
    assert (second.passport_series, second.passport_number) == ("6502", "654321")
    assert second.passport_issuer == "УФМС России"
    assert second.address == "г. Екатеринбург, ул. Мира, д. 2"


def test_parser_splits_grouped_passport_headings_and_applies_shared_fields(tmp_path):
    path = tmp_path / "grouped-passport-headings.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Платформа скоринга"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада Автора 1",
            "Паспортные данные Автора 1*\n"
            "- Иванов Иван Иванович;\n"
            "- 01.01.1990;\n"
            "- Россия, г. Екатеринбург, ул. Ленина, д. 1;\n"
            "- серия 6501 и номер 123456 паспорта, дата выдачи 01.02.2010 МВД России\n"
            "Паспортные данные Автора 2*\n"
            "- Петров Петр Петрович;\n"
            "- 02.02.1992;\n"
            "- Россия, г. Екатеринбург, ул. Мира, д. 2;\n"
            "- серия 6502 и номер 654321 паспорта, дата выдачи 03.04.2012 УФМС России",
        ),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада Автора 1",
            "- написание исходного текста программы;\n- разработка алгоритма",
        ),
        (
            "Основание возникновения права (указывается для каждого автора)",
            "- заявитель является работодателем автора",
        ),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert len(result.data.authors) == 2
    first, second = result.data.authors
    assert (first.full_name, first.passport_series, first.passport_number) == (
        "Иванов Иван Иванович", "6501", "123456"
    )
    assert (second.full_name, second.passport_series, second.passport_number) == (
        "Петров Петр Петрович", "6502", "654321"
    )
    assert first.creative_contribution == second.creative_contribution == (
        "написание исходного текста программы; разработка алгоритма"
    )
    assert first.rights_basis == second.rights_basis == (
        "заявитель является работодателем автора"
    )


def test_numbered_contributions_fill_existing_grouped_authors(tmp_path):
    path = tmp_path / "grouped-authors-separate-contributions.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Террамаркет"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Количество авторов", "2"),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада Автора 1",
            "Авторы:\n"
            "1. Иванов Иван Иванович\nПаспорт 6501 123456 выдан МВД России\n"
            "Дата выдачи 01.02.2010\nМесто жительства: г. Москва, д. 1\n"
            "Дата рождения: 01.01.1990\n"
            "2. Петров Петр Петрович\nПаспорт 6502 654321 выдан УФМС России\n"
            "Дата выдачи 03.04.2012\nМесто жительства: г. Москва, д. 2\n"
            "Дата рождения: 02.02.1992",
        ),
        (
            "Паспортные данные Автора 1 / Описание творческого вклада Автора 1",
            "1. Иванов Иван Иванович\n- разработка алгоритма\n"
            "2. Петров Петр Петрович\n- написание исходного кода",
        ),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert len(result.data.authors) == 2
    first, second = result.data.authors
    assert first.creative_contribution == "разработка алгоритма"
    assert second.creative_contribution == "написание исходного кода"
    assert first.address == "г. Москва, д. 1"
    assert second.address == "г. Москва, д. 2"


def test_parser_supports_colons_after_passport_series_and_issued_labels(tmp_path):
    path = tmp_path / "colon-separated-passport.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Программа управления оборудованием"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1",
            "Иванов Кирилл Олегович,\n"
            "дата рождения: 22.07.2000\n"
            "паспорт серия: 80 25 № 110704,\n"
            "дата выдачи - 05.08.2020,\n"
            "выдан: МВД ПО РЕСПУБЛИКЕ БАШКОРТОСТАН,\n"
            "зарегистрирован: 450000, Республика Башкортостан, г. Уфа, ул. Ленина, д. 1",
        ),
        ("Описание творческого вклада автора 1", "Разработка программы"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.passport_series == "8025"
    assert author.passport_number == "110704"
    assert author.passport_issue_date == "05.08.2020"
    assert author.passport_issuer == "МВД ПО РЕСПУБЛИКЕ БАШКОРТОСТАН"
    assert author.address == (
        "450000, Республика Башкортостан, г. Уфа, ул. Ленина, д. 1"
    )


def test_parser_supports_year_first_text_birth_date(tmp_path):
    path = tmp_path / "year-first-birth-date.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Наименование программы", "Программа управления оборудованием"),
        ("Программа создана за счет средств:", "Собственных"),
        ("Наименование заявителя и ИНН", "ООО «Пример», ИНН 7707083893"),
        ("Количество авторов", "1"),
        (
            "Паспортные данные Автора 1",
            "- Петров Сергей Вячеславович;\n"
            "- 1982 г. 3 января\n"
            "- г. Екатеринбург, ул. Гражданская, д. 11, кв. 29\n"
            "- Паспорт: серия 6520, номер 187894; "
            "Дата выдачи 24.11.2020; "
            "Выдан ГУ МВД России по Свердловской области",
        ),
        ("Описание творческого вклада автора 1", "Разработка программы"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    author = parse_questionnaire(path).data.authors[0]

    assert author.birth_date == "03.01.1982"
    assert author.address == "г. Екатеринбург, ул. Гражданская, д. 11, кв. 29"
    assert author.passport_issue_date == "24.11.2020"


def test_parser_does_not_warn_about_passport_data_when_authors_are_not_mentioned(tmp_path):
    path = tmp_path / "authors-not-mentioned.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("РАЗДЕЛ I – СВЕДЕНИЯ О ПО ДЛЯ РОСПАТЕНТА", ""),
        ("Наименование программы", "Тестовая программа"),
        ("Наименование заявителя и ИНН", "ООО «Тест», ИНН 7707083893"),
        ("Количество авторов", "1"),
        ("Будут ли упоминаться авторы при регистрации программы для ЭВМ?", "Нет"),
        ("ФИО авторов", "Иванов Иван Иванович"),
        ("Творческий вклад авторов", "Автор 1:\n- Разработка алгоритма"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    assert result.data.authors_will_be_mentioned is False
    author_warnings = [warning for warning in result.warnings if warning.startswith("Автор 1:")]
    assert author_warnings == []


def test_parser_keeps_non_passport_warnings_when_authors_are_not_mentioned(tmp_path):
    path = tmp_path / "authors-not-mentioned-and-empty.docx"
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("РАЗДЕЛ I – СВЕДЕНИЯ О ПО ДЛЯ РОСПАТЕНТА", ""),
        ("Наименование программы", "Тестовая программа"),
        ("Наименование заявителя и ИНН", "ООО «Тест», ИНН 7707083893"),
        ("Количество авторов", "1"),
        ("Будут ли указываться авторы при регистрации программы для ЭВМ?", "нет"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(path)

    result = parse_questionnaire(path)

    author_warnings = [warning for warning in result.warnings if warning.startswith("Автор 1:")]
    assert author_warnings == ["Автор 1: не найдены ФИО, творческий вклад."]
